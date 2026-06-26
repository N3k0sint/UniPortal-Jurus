#!/usr/bin/env python3
"""
Generate Jurus_FarhanIsmail.docx from the markdown report content.
Uses python-docx to create a professional Word document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helper Functions ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p

def add_code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_n = p.add_run(text)
        run_n.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('JURUS LEVEL 1 ANALYST')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('TECHNICAL REPORT')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

proj = doc.add_paragraph()
proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = proj.add_run('UniKL PJ Research Collaboration Portal')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
doc.add_paragraph()

for line in [
    'Prepared by: Farhan Ismail',
    'Programme: JURUS – Jurutera Siber RAWSEC',
    'Level: Level 1 Analyst (Foundational Operator)',
    'Challenge: Cyber Engineering Challenge',
    'Date Submitted: 20 June 2026',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(60, 60, 60)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════
add_heading_styled('TABLE OF CONTENTS', level=1)
toc_items = [
    '1. Executive Summary & Architecture',
    '2. Module 1 – Operating System Engineering',
    '3. Module 2 – Network Security',
    '4. Module 3 – Database Security',
    '5. Module 4 – Application Security',
    '6. Module 5 – Security Monitoring',
    '7. Module 6 – Business Continuity Planning (BCP/DR)',
    '8. Conclusion',
    '9. References',
    'Appendix A – Architecture Diagram',
    'Appendix B – Full Script Listings',
]
for item in toc_items:
    add_body(item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════
add_heading_styled('1. EXECUTIVE SUMMARY & ARCHITECTURE', level=1)

add_heading_styled('1.1 Project Overview', level=2)
add_body('This report documents the design, implementation, security hardening, monitoring, and disaster recovery of a production-ready Research Collaboration Portal for the UniKL Petaling Jaya (PJ) University Consortium. The portal enables university researchers to manage collaborative research projects, securely upload and share academic documents, and track collaboration activities through a role-based access control (RBAC) system.')
add_body('The entire solution was independently designed, deployed, secured, and documented within a 3-week challenge period, following an engineering-first philosophy where no predefined templates, reference architectures, or vendor-specific technologies were provided.')

add_heading_styled('1.2 Environment Specification', level=2)
add_table(
    ['Component', 'Specification'],
    [
        ['Operating System', 'Linux Mint 22 (Ubuntu 24.04 LTS base)'],
        ['Virtualisation', 'VMware Workstation Pro (Single-Host VM)'],
        ['Web Framework', 'Python Flask 3.0.3'],
        ['Application Server', 'Gunicorn 23.0.0 (3 workers, WSGI)'],
        ['Reverse Proxy / TLS', 'Nginx (HTTPS, TLS 1.3, Self-Signed Certificate)'],
        ['Database', 'PostgreSQL 16 (Production)'],
        ['Active Defense', 'Fail2ban (SSH + Nginx jails)'],
        ['Host Firewall', 'UFW (Uncomplicated Firewall)'],
        ['Log Management', 'Rsyslog + Logwatch + Logrotate'],
        ['Backup Encryption', 'GPG Symmetric AES-256'],
        ['Version Control', 'Git + GitHub (Private Repository)'],
        ['Security Scanning', 'Snyk Code, GitHub CodeQL'],
    ]
)

add_heading_styled('1.3 Technology Selection Rationale', level=2)
add_bullet('Linux Mint / Ubuntu 24.04 LTS was chosen for its long-term support (LTS), robust security update pipeline, and strong community documentation. Using a single-host VM was a practical decision to maintain full control within the 3-week time constraint.', bold_prefix='Linux Mint / Ubuntu 24.04 LTS – ')
add_bullet('Flask was chosen over Django for its lightweight modularity — it allowed precise control over every security component (CSRF, session management, rate limiting) without hidden abstractions.', bold_prefix='Flask – ')
add_bullet('PostgreSQL was chosen over MariaDB/MySQL for its superior built-in security controls (pg_hba.conf host access control, schema-level privilege grants, and native pg_dump/pg_restore for BCP).', bold_prefix='PostgreSQL – ')
add_bullet('Nginx was chosen as a reverse proxy to enforce TLS 1.3, inject security headers at the edge, and separate static file serving from the application server — reducing attack surface on Gunicorn.', bold_prefix='Nginx – ')

add_heading_styled('1.4 Logical Architecture Diagram', level=2)
add_body('(Refer to the architecture diagram in Appendix A or Section 1.4 of the Markdown report. A visual diagram should be inserted here using draw.io or similar tool.)', italic=True)

add_code_block("""INTERNET (Browser HTTPS:443)
        │
  ┌─────▼─────┐
  │ UFW (deny) │  Allow: 80, 443, 2222
  └─────┬─────┘
  ┌─────▼─────┐
  │   Nginx    │  TLS 1.3, Security Headers, server_tokens off
  └─────┬─────┘
        │ proxy_pass http://127.0.0.1:5000
  ┌─────▼─────┐
  │  Gunicorn  │  3 Workers, Bind: localhost only
  └─────┬─────┘
  ┌─────▼─────┐
  │ Flask App  │  CSRF, RBAC, Rate Limiting, Audit Logging
  └──┬────┬───┘
     │    │
  PostgreSQL   /var/www/uniklpj_uploads/
  (localhost)  (chmod 770, PDF storage)

  Fail2ban │ Rsyslog │ Cron (backup.sh daily 02:00)""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. MODULE 1 – OS ENGINEERING
# ═══════════════════════════════════════════════════════════════
add_heading_styled('2. MODULE 1 – OPERATING SYSTEM ENGINEERING', level=1)

add_heading_styled('2.1 Objective', level=2)
add_body('Build a secure host baseline on Linux by hardening user access controls, enforcing strong password policies, disabling unnecessary services, and configuring a host-level firewall.')

add_heading_styled('2.2 User Privilege Management & Sudo Restriction', level=2)
add_body('Root login is fully disabled. The system operator account "jurus" is the only account with controlled sudo privileges. No direct root login is permitted via SSH or console.')
add_body('Evidence – SSH Daemon Configuration (/etc/ssh/sshd_config):', bold=True)
add_code_block("""Port 2222
PermitRootLogin no
PasswordAuthentication no
PubkeyAuthentication yes
MaxAuthTries 3
X11Forwarding no""")

add_body('Hardening Applied:', bold=True)
add_bullet('PermitRootLogin no — Eliminates root SSH brute-force attacks entirely.', bold_prefix='')
add_bullet('PasswordAuthentication no — Enforces SSH key-based authentication only.', bold_prefix='')
add_bullet('Port 2222 — Default SSH port changed to reduce automated scanning noise.', bold_prefix='')
add_bullet('MaxAuthTries 3 — Limits authentication attempts per connection to 3.', bold_prefix='')
add_bullet('X11Forwarding no — Disables GUI forwarding to eliminate X11 attack surface.', bold_prefix='')

add_body('Script Implementation (from setup_system.sh):', bold=True)
add_code_block("""SSH_CONFIG="/etc/ssh/sshd_config"
cp "$SSH_CONFIG" "${SSH_CONFIG}.bak"
sed -i 's/^#\\?Port.*/Port 2222/' "$SSH_CONFIG"
sed -i 's/^#\\?PermitRootLogin.*/PermitRootLogin no/' "$SSH_CONFIG"
sed -i 's/^#\\?PasswordAuthentication.*/PasswordAuthentication no/' "$SSH_CONFIG"
sed -i 's/^#\\?PubkeyAuthentication.*/PubkeyAuthentication yes/' "$SSH_CONFIG"
sed -i 's/^#\\?MaxAuthTries.*/MaxAuthTries 3/' "$SSH_CONFIG"
sed -i 's/^#\\?X11Forwarding.*/X11Forwarding no/' "$SSH_CONFIG"
systemctl restart sshd || systemctl restart ssh""")

add_heading_styled('2.3 Password Complexity Policy (PAM)', level=2)
add_code_block('password requisite pam_pwquality.so retry=3 minlen=8 ucredit=-1 lcredit=-1 dcredit=-1 ocredit=-1')

add_table(
    ['Parameter', 'Requirement'],
    [
        ['minlen=8', 'Minimum 8 characters'],
        ['ucredit=-1', 'At least 1 uppercase letter'],
        ['lcredit=-1', 'At least 1 lowercase letter'],
        ['dcredit=-1', 'At least 1 digit'],
        ['ocredit=-1', 'At least 1 special character'],
        ['retry=3', 'Maximum 3 retries on failure'],
    ]
)

add_heading_styled('2.4 Host Firewall Configuration (UFW)', level=2)
add_code_block("""ufw --force reset
ufw default deny incoming
ufw default allow outgoing
ufw allow 2222/tcp comment 'Hardened SSH'
ufw allow 80/tcp comment 'HTTP Redirect'
ufw allow 443/tcp comment 'HTTPS Portal'
ufw --force enable""")

add_table(
    ['Rule', 'Direction', 'Port', 'Protocol', 'Purpose'],
    [
        ['ALLOW', 'IN', '2222', 'TCP', 'Hardened SSH access'],
        ['ALLOW', 'IN', '80', 'TCP', 'HTTP → HTTPS 301 redirect'],
        ['ALLOW', 'IN', '443', 'TCP', 'HTTPS portal access'],
        ['DENY', 'IN', '*', '*', 'All other ports blocked (default)'],
    ]
)

add_heading_styled('2.5 Service Minimisation', level=2)
add_body('All unnecessary services have been disabled. Only essential network services are active:')
add_table(
    ['Service', 'Port', 'Binding', 'Purpose'],
    [
        ['SSH (sshd)', '2222', '0.0.0.0', 'Remote admin (key-only)'],
        ['Nginx', '80, 443', '0.0.0.0', 'Web reverse proxy'],
        ['Gunicorn', '5000', '127.0.0.1', 'Flask app (localhost only)'],
        ['PostgreSQL', '5432', '127.0.0.1', 'Database (localhost only)'],
    ]
)
add_body('Note: Both Gunicorn and PostgreSQL are bound to 127.0.0.1 (localhost) only and are not exposed to external networks.')

add_heading_styled('2.6 Directory & File Permissions', level=2)
add_code_block("""chown -R jurus:www-data /var/www/uniklpj_uploads    # chmod 770
chown -R jurus:www-data /var/log/uniklpj            # chmod 770
chown -R jurus:www-data /var/www/uniklpj/static     # chmod 750
chmod 600 /etc/ssl/private/uniklpj.key              # Root-only""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. MODULE 2 – NETWORK SECURITY
# ═══════════════════════════════════════════════════════════════
add_heading_styled('3. MODULE 2 – NETWORK SECURITY', level=1)

add_heading_styled('3.1 Nginx Reverse Proxy & TLS 1.3 Enforcement', level=2)
add_body('All client traffic is encrypted using TLS 1.3. Legacy protocols (TLS 1.0, 1.1, 1.2) are explicitly disabled to prevent downgrade attacks.')
add_body('Full Nginx Configuration (nginx_uniklpj.conf):', bold=True)
add_code_block("""server {
    listen 80;
    server_name localhost;
    server_tokens off;
    return 301 https://$host$request_uri;
}

server {
    listen 443 ssl http2;
    server_name localhost;

    ssl_certificate /etc/ssl/certs/uniklpj.crt;
    ssl_certificate_key /etc/ssl/private/uniklpj.key;
    ssl_protocols TLSv1.3;
    ssl_prefer_server_ciphers off;
    server_tokens off;
    client_max_body_size 10M;

    add_header X-Frame-Options "SAMEORIGIN" always;
    add_header X-Content-Type-Options "nosniff" always;
    add_header Content-Security-Policy "default-src 'self'; ..." always;
    add_header Referrer-Policy "strict-origin-when-cross-origin" always;
    add_header Strict-Transport-Security "max-age=31536000; includeSubDomains" always;

    location / {
        proxy_pass http://127.0.0.1:5000;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
    }
    location /static/ {
        alias /var/www/uniklpj/static/;
        expires 30d;
    }
}""")

add_heading_styled('3.2 HTTP Security Headers', level=2)
add_table(
    ['Header', 'Value', 'Purpose'],
    [
        ['X-Frame-Options', 'SAMEORIGIN', 'Prevents clickjacking'],
        ['X-Content-Type-Options', 'nosniff', 'Prevents MIME-type sniffing'],
        ['Content-Security-Policy', "default-src 'self'; ...", 'Restricts resources to same-origin, blocks XSS'],
        ['Referrer-Policy', 'strict-origin-when-cross-origin', 'Limits referrer leakage'],
        ['Strict-Transport-Security', 'max-age=31536000; includeSubDomains', 'Enforces HTTPS for 1 year (HSTS)'],
    ]
)

add_heading_styled('3.3 Version Banner Suppression', level=2)
add_body('The server_tokens off directive removes the Nginx version number from all HTTP response headers and error pages. Attackers cannot fingerprint the server software version for targeted exploit searches.')

add_heading_styled('3.4 Self-Signed SSL Certificate Generation', level=2)
add_code_block("""openssl req -x509 -nodes -days 365 -newkey rsa:2048 \\
  -keyout /etc/ssl/private/uniklpj.key \\
  -out /etc/ssl/certs/uniklpj.crt \\
  -subj "/C=MY/ST=Selangor/L=Petaling Jaya/O=UniKL PJ/OU=Cybersecurity/CN=localhost"
chmod 600 /etc/ssl/private/uniklpj.key""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. MODULE 3 – DATABASE SECURITY
# ═══════════════════════════════════════════════════════════════
add_heading_styled('4. MODULE 3 – DATABASE SECURITY', level=1)

add_heading_styled('4.1 PostgreSQL Deployment & Least Privilege', level=2)
add_body('A dedicated, non-superuser PostgreSQL account "portal_user" is created with access limited to only the uniklpj_db database:')
add_code_block("""sudo -u postgres psql -c "CREATE DATABASE uniklpj_db;"
sudo -u postgres psql -c "CREATE USER portal_user WITH ENCRYPTED PASSWORD '***';"
sudo -u postgres psql -c "GRANT ALL PRIVILEGES ON DATABASE uniklpj_db TO portal_user;"
sudo -u postgres psql -d uniklpj_db -c "GRANT ALL ON SCHEMA public TO portal_user;" """)

add_table(
    ['Property', 'Value'],
    [
        ['Username', 'portal_user'],
        ['Superuser?', 'NO'],
        ['Create DB?', 'NO'],
        ['Create Role?', 'NO'],
        ['Accessible Databases', 'uniklpj_db only'],
        ['Schema Access', 'public schema only'],
    ]
)

add_heading_styled('4.2 Host Access Restriction (pg_hba.conf)', level=2)
add_code_block("""# TYPE  DATABASE        USER            ADDRESS         METHOD
local   all             postgres                        peer
host    uniklpj_db      portal_user     127.0.0.1/32    md5
host    uniklpj_db      portal_user     ::1/128         md5""")
add_bullet('127.0.0.1/32 — Only accepts connections from localhost. No remote access.', bold_prefix='')
add_bullet('md5 — Password authentication uses MD5 hash challenge. Never plaintext.', bold_prefix='')
add_bullet('postgres superuser uses peer authentication (Unix socket only).', bold_prefix='')

add_heading_styled('4.3 Secure Credential Handling', level=2)
add_body('Database credentials are stored in an environment file (.env) that is excluded from version control via .gitignore. The .env.example template is committed as a safe reference.')
add_code_block("""# .gitignore
.env
instance/
*.db""")

add_heading_styled('4.4 Database Schema Design', level=2)
add_table(
    ['Table', 'Purpose', 'Key Columns'],
    [
        ['roles', 'RBAC role definitions', 'id, role_name'],
        ['users', 'User accounts with hashed passwords', 'id, username, email, password_hash, role_id (FK), is_blocked'],
        ['projects', 'Research projects', 'id, title, description, owner_id (FK)'],
        ['documents', 'Uploaded PDF documents', 'id, original_filename, secure_uuid_filename, mime_type, is_approved, delete_requested'],
        ['audit_logs', 'Security event trail', 'id, user_id (FK), action, ip_address, details, timestamp'],
        ['blocked_ips', 'Dynamic IP blacklist', 'id, ip_address, reason, blocked_at'],
        ['collaborations', 'Project collaboration links', 'id, project_id (FK), collaborator_id (FK), status'],
    ]
)
add_body('Password Storage: All passwords are hashed using bcrypt (via Flask-Bcrypt). Plaintext passwords are never stored in the database.', bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. MODULE 4 – APPLICATION SECURITY
# ═══════════════════════════════════════════════════════════════
add_heading_styled('5. MODULE 4 – APPLICATION SECURITY', level=1)

add_heading_styled('5.1 Authentication & Session Management', level=2)
add_code_block("""app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SECURE'] = True  # Production
app.config['SESSION_COOKIE_SAMESITE'] = 'Strict'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=1)
login_manager.session_protection = "strong" """)

add_table(
    ['Control', 'Implementation'],
    [
        ['HttpOnly cookies', 'Prevents JavaScript access to session cookies'],
        ['Secure cookies', 'Cookies only transmitted over HTTPS'],
        ['SameSite', 'Prevents CSRF via cookie scope restriction'],
        ['Session timeout', '1-hour inactivity timeout'],
        ['Anti-hijacking', 'Invalidates sessions if IP/user-agent changes'],
    ]
)

add_heading_styled('5.2 CSRF Protection', level=2)
add_body('Cross-Site Request Forgery protection is enabled globally using Flask-WTF CSRFProtect. Every HTML form includes a hidden CSRF token. All POST requests without a valid token are rejected with a 400 error.')

add_heading_styled('5.3 Input Validation (WTForms)', level=2)
add_body('All user inputs are validated using WTForms validators with strict whitelisting regex patterns. SQL Injection is prevented by using Flask-SQLAlchemy ORM exclusively — no raw SQL queries exist in the codebase.')
add_code_block("""username = StringField('Username', validators=[
    DataRequired(), Length(min=3, max=25),
    Regexp(r'^[a-zA-Z0-9_]+$', message="Letters, numbers, underscores only.")
])

password = PasswordField('Password', validators=[
    DataRequired(), Length(min=8, max=128),
    Regexp(r'^(?=.*[a-z])(?=.*[A-Z])(?=.*\\d)(?=.*[@$!%*?&])',
           message="Must include uppercase, lowercase, digit, special character.")
])""")

add_heading_styled('5.4 Role-Based Access Control (RBAC)', level=2)
add_table(
    ['Role', 'Privileges'],
    [
        ['Admin', 'Full system access: manage users, view audit logs, block/unblock IPs, block/suspend users, generate anomaly reports'],
        ['Researcher', 'Create projects, upload documents, manage collaborators, approve/reject uploads and deletion requests'],
        ['Collaborator', 'Request collaboration, upload documents (pending approval), request document deletions'],
    ]
)

add_heading_styled('5.5 Secure File Upload', level=2)
add_table(
    ['Control', 'Implementation'],
    [
        ['Extension Whitelist', 'Only .pdf files allowed'],
        ['MIME Type Verification', 'python-magic verifies actual file signature (magic bytes)'],
        ['UUID Renaming', 'Files renamed to random UUIDs to prevent path traversal'],
        ['Storage Outside Web Root', 'Files stored in /var/www/uniklpj_uploads/ (not in app directory)'],
        ['Size Limit', '10MB enforced at both Flask and Nginx levels'],
        ['Approval Workflow', 'Collaborator uploads require owner approval'],
    ]
)

add_heading_styled('5.6 Open Redirect Mitigation', level=2)
add_body('All redirect(request.url) calls were replaced with safe, validated redirect URLs using url_for() to prevent CWE-601 Open Redirect vulnerabilities (flagged by Snyk Code and GitHub CodeQL).')

add_heading_styled('5.7 Rate Limiting (DDoS / Brute-Force Protection)', level=2)
add_body('Application-level rate limiting is implemented using Flask-Limiter with a secure IP extraction function that handles X-Forwarded-For proxy chains:')
add_table(
    ['Endpoint', 'Rate Limit', 'Purpose'],
    [
        ['/login (POST)', '5 per minute', 'Brute-force login protection'],
        ['/register (POST)', '10 per hour', 'Registration spam prevention'],
        ['/project/new (POST)', '10 per hour', 'Project creation abuse prevention'],
        ['/project/<id>/upload (POST)', '10 per minute', 'File upload flood prevention'],
    ]
)

add_heading_styled('5.8 Custom Error Pages', level=2)
add_table(
    ['Error Code', 'Template', 'Description'],
    [
        ['400', 'errors/400.html', 'Bad Request'],
        ['403', 'errors/403.html', 'Forbidden / IP Blocked'],
        ['404', 'errors/404.html', 'Not Found'],
        ['429', 'errors/429.html', 'Rate Limited (Too Many Requests)'],
        ['500', 'errors/500.html', 'Internal Server Error (no stack trace exposed)'],
    ]
)

add_heading_styled('5.9 Dependency Security', level=2)
add_body('All dependencies are pinned to specific, secure versions. Werkzeug was upgraded from 2.2.3 to 3.1.8 to fix CVE-2024-34069 (Remote Code Execution, CVSS 7.5). Snyk Code and GitHub CodeQL were used for SAST scanning.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. MODULE 5 – SECURITY MONITORING
# ═══════════════════════════════════════════════════════════════
add_heading_styled('6. MODULE 5 – SECURITY MONITORING', level=1)

add_heading_styled('6.1 Application-Level Audit Logging', level=2)
add_body('Every security-relevant action is logged to the audit_logs database table with full context:')
add_table(
    ['Log Event', 'Trigger'],
    [
        ['USER_LOGIN_SUCCESS', 'Successful authentication'],
        ['USER_LOGIN_FAILED', 'Failed login attempt (wrong credentials)'],
        ['USER_LOGIN_BLOCKED', 'Login attempt on suspended/blocked account'],
        ['USER_REGISTRATION', 'New user registration'],
        ['USER_LOGOUT', 'User sign-out'],
        ['PROJECT_CREATED', 'New project creation'],
        ['DOCUMENT_UPLOADED', 'PDF document upload'],
        ['DOCUMENT_DELETED', 'Document deletion'],
        ['MALICIOUS_FILE_UPLOAD_BLOCKED', 'Upload blocked (invalid extension)'],
        ['MALICIOUS_FILE_SIGNATURE_BLOCKED', 'Upload blocked (MIME mismatch)'],
        ['UNAUTHORIZED_ACCESS_ATTEMPT', 'RBAC violation attempt'],
        ['RATE_LIMIT_EXCEEDED', 'Rate limit threshold exceeded'],
        ['COLLABORATION_REQUESTED', 'Collaboration request submitted'],
        ['COLLABORATION_APPROVED/REJECTED', 'Collaboration decision'],
        ['USER_BLOCKED / USER_UNBLOCKED', 'Admin suspended/reactivated account'],
        ['IP_BLOCKED / IP_UNBLOCKED', 'Admin blocked/unblocked IP address'],
    ]
)
add_body('Each log entry records: User ID, Action type, Client IP address (with X-Forwarded-For proxy resolution), Detailed description, and Timestamp (Malaysia Time, UTC+8).')

add_heading_styled('6.2 Admin Audit Log Viewer & Anomaly Reports', level=2)
add_body('Administrators can view and search the full audit trail directly from the admin dashboard (/admin/logs). The admin panel also includes a PDF Anomaly Report generator that compiles security metrics including failed login attempts, blocked accounts, blocked IPs, rate limit violations, and file upload security blocks.')

add_heading_styled('6.3 Fail2ban Active Defense', level=2)
add_code_block("""[sshd]
enabled = true
port = 2222
filter = sshd
logpath = /var/log/auth.log
maxretry = 5
bantime = 1800
findtime = 600

[nginx-http-auth]
enabled = true
filter = nginx-http-auth
port = http,https
logpath = /var/log/nginx/error.log
maxretry = 5
bantime = 1800
findtime = 600""")

add_table(
    ['Jail', 'Trigger', 'Ban Duration', 'Find Window'],
    [
        ['sshd', '5 failed SSH logins', '30 minutes (1800s)', '10 minutes (600s)'],
        ['nginx-http-auth', '5 failed HTTP auth attempts', '30 minutes (1800s)', '10 minutes (600s)'],
    ]
)

add_heading_styled('6.4 Gunicorn Access & Error Logs', level=2)
add_code_block("""ExecStart=.../gunicorn --workers 3 --bind 127.0.0.1:5000 \\
    --access-logfile /var/log/uniklpj/access.log \\
    --error-logfile /var/log/uniklpj/error.log \\
    wsgi:app""")

add_heading_styled('6.5 Log Rotation (logrotate)', level=2)
add_code_block("""/var/log/uniklpj/*.log {
    daily
    rotate 7
    compress
    delaycompress
    missingok
    notifempty
    create 0660 jurus www-data
}""")

add_heading_styled('6.6 System-Level Logging', level=2)
add_bullet('Rsyslog collects and centralises all system logs including kernel messages, authentication events, and service status changes.')
add_bullet('Logwatch provides daily summaries of system security events including SSH access, service failures, and disk usage alerts.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. MODULE 6 – BCP/DR
# ═══════════════════════════════════════════════════════════════
add_heading_styled('7. MODULE 6 – BUSINESS CONTINUITY PLANNING (BCP/DR)', level=1)

add_heading_styled('7.1 BCP/DR Parameters', level=2)
add_table(
    ['Parameter', 'Full Name', 'Value', 'Justification'],
    [
        ['MTD', 'Maximum Tolerable Downtime', '4 hours', 'Research portal users can tolerate up to 4 hours of downtime before academic workflows are critically impacted.'],
        ['MTO', 'Maximum Tolerable Outage', '2 hours', 'Portal restoration must begin within 2 hours of incident detection to remain within the MTD window.'],
        ['RTO', 'Recovery Time Objective', '< 5 minutes', 'The automated restore.sh script completes full recovery in ~15-30 seconds. Including manual passphrase entry, recovery is within 5 minutes.'],
        ['RPO', 'Recovery Point Objective', '24 hours', 'Automated daily backups via Cron Job at 02:00 MYT. Maximum data loss in worst-case is 24 hours of changes.'],
    ]
)

add_heading_styled('7.2 BCP/DR Workflow', level=2)
add_body('Step 1 – BACKUP (Daily Automated via Cron):', bold=True)
add_bullet('pg_dump → PostgreSQL database dump')
add_bullet('cp -r → Upload files copied')
add_bullet('tar -czf → Archive compressed')
add_bullet('gpg --symmetric --cipher-algo AES256 → Encrypted')
add_bullet('chmod 600 → Secure file permissions')
add_bullet('find -mtime +7 -delete → 7-day retention cleanup')
add_body('')
add_body('Step 2 – SIMULATED CRASH:', bold=True)
add_bullet('systemctl stop uniklpj → Stop web application')
add_bullet('DROP DATABASE uniklpj_db → Destroy the database')
add_bullet('rm -rf /var/www/uniklpj_uploads/* → Wipe all uploaded files')
add_body('')
add_body('Step 3 – DISASTER RECOVERY (RTO Measured):', bold=True)
add_bullet('START_TIME=$(date +%s) ← RTO Stopwatch begins')
add_bullet('gpg --decrypt → Decrypt latest .gpg backup')
add_bullet('tar -xzf → Extract archive')
add_bullet('Restore uploaded files and database (pg_restore)')
add_bullet('systemctl start uniklpj → Bring portal online')
add_bullet('ELAPSED=$((END_TIME - START_TIME)) → Report RTO')

add_heading_styled('7.3 Backup Script (scripts/backup.sh)', level=2)
add_code_block("""#!/bin/bash
# UniKL PJ Collaboration Portal - BCP Encryption Backup Script (RPO)
BACKUP_DIR="/var/backups/uniklpj"
UPLOAD_DIR="/var/www/uniklpj_uploads"
DB_NAME="uniklpj_db"
DB_USER="portal_user"
TIMESTAMP=$(date +"%Y%m%d_%H%M%S")
BACKUP_NAME="uniklpj_backup_${TIMESTAMP}"
PASSPHRASE=${BACKUP_PASSPHRASE:-"UniKLBackupPassphrase2026!"}

mkdir -p "$BACKUP_DIR" && chmod 700 "$BACKUP_DIR"
TEMP_DIR="/tmp/uniklpj_backup_tmp"
rm -rf "$TEMP_DIR" && mkdir -p "$TEMP_DIR"

# 1. PostgreSQL Dump
export PGPASSWORD="PortalSecure@2026!"
pg_dump -h localhost -U "$DB_USER" -F c -b -f "$TEMP_DIR/db_dump.dump" "$DB_NAME"

# 2. Archive
cp -r "$UPLOAD_DIR" "$TEMP_DIR/uploads"
tar -czf "$TEMP_DIR/archive.tar.gz" -C "$TEMP_DIR" db_dump.dump uploads

# 3. Encrypt (AES-256)
gpg --batch --yes --passphrase "$PASSPHRASE" --symmetric \\
    --cipher-algo AES256 -o "$BACKUP_DIR/${BACKUP_NAME}.tar.gz.gpg" \\
    "$TEMP_DIR/archive.tar.gz"
chmod 600 "$BACKUP_DIR/${BACKUP_NAME}.tar.gz.gpg"
rm -rf "$TEMP_DIR"

# 4. Retention Policy (7 days)
find "$BACKUP_DIR" -name "uniklpj_backup_*.tar.gz.gpg" -mtime +7 -delete""")

add_heading_styled('7.4 Crash Simulation Script (scripts/simulate_crash.sh)', level=2)
add_code_block("""#!/bin/bash
# Disaster Crash Simulation Utility
systemctl stop uniklpj
sudo -u postgres psql -c "DROP DATABASE IF EXISTS uniklpj_db;"
rm -rf /var/www/uniklpj_uploads/*
systemctl start uniklpj
echo "[!] CRASH SIMULATION COMPLETE. Run 'sudo ./restore.sh' to recover." """)

add_heading_styled('7.5 Disaster Recovery Script (scripts/restore.sh)', level=2)
add_code_block("""#!/bin/bash
# Disaster Recovery Restore Script (RTO)
LATEST_BACKUP=$(ls -t /var/backups/uniklpj/uniklpj_backup_*.tar.gz.gpg | head -1)
read -s -p "Enter GPG Passphrase: " PASSPHRASE
START_TIME=$(date +%s)                          # RTO Stopwatch Start

systemctl stop uniklpj

# 1. Decrypt
gpg --batch --passphrase "$PASSPHRASE" --decrypt -o /tmp/archive.tar.gz "$LATEST_BACKUP"
# 2. Extract
tar -xzf /tmp/archive.tar.gz -C /tmp/uniklpj_restore_tmp
# 3. Restore files
cp -r /tmp/uniklpj_restore_tmp/uploads /var/www/uniklpj_uploads
chown -R jurus:www-data /var/www/uniklpj_uploads && chmod 770 /var/www/uniklpj_uploads
# 4. Restore database
sudo -u postgres psql -c "DROP DATABASE IF EXISTS uniklpj_db;"
sudo -u postgres psql -c "CREATE DATABASE uniklpj_db OWNER portal_user;"
pg_restore -h localhost -U portal_user -d uniklpj_db /tmp/uniklpj_restore_tmp/db_dump.dump
# 5. Bring online
systemctl start uniklpj

END_TIME=$(date +%s)
ELAPSED=$((END_TIME - START_TIME))              # RTO Stopwatch End
echo "[+] Total Recovery Time (RTO): ${ELAPSED} seconds." """)

add_heading_styled('7.6 Cron Job (Automated Daily Backup)', level=2)
add_code_block("""# Crontab entry (sudo crontab -e)
0 2 * * * /home/jurus/Documents/Project/UniPortal-Jurus/scripts/backup.sh >> /var/log/uniklpj/backup.log 2>&1""")
add_body('This executes backup.sh every day at 02:00 MYT, ensuring the RPO of 24 hours is consistently met.')

add_heading_styled('7.7 BCP Parameter Compliance Summary', level=2)
add_table(
    ['Parameter', 'Target', 'Achieved', 'Evidence'],
    [
        ['RPO', '≤ 24 hours', '✓ 24 hours', 'Daily Cron backup at 02:00 MYT'],
        ['RTO', '< 5 minutes', '✓ ~15-30 seconds', 'restore.sh reports elapsed time'],
        ['MTO', '≤ 2 hours', '✓ Achievable', 'Single command + passphrase'],
        ['MTD', '≤ 4 hours', '✓ Within margin', 'Total recovery < 30 minutes'],
        ['Encryption', 'AES-256', '✓', 'GPG symmetric AES256'],
        ['Retention', '7 days', '✓', 'find -mtime +7 -delete'],
        ['Integrity', 'Secure permissions', '✓', 'chmod 600 (files), chmod 700 (dir)'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. CONCLUSION
# ═══════════════════════════════════════════════════════════════
add_heading_styled('8. CONCLUSION', level=1)
add_body('This report documents the end-to-end design, implementation, security hardening, monitoring, and disaster recovery of the UniKL PJ Research Collaboration Portal across all six JURUS modules:')
add_body('')
add_bullet('OS Engineering — Root login disabled, SSH key-only authentication on non-standard port 2222, PAM password complexity enforcement, UFW firewall with default-deny policy, and service minimisation.', bold_prefix='Module 1: ')
add_bullet('Network Security — Nginx reverse proxy with TLS 1.3 enforcement (legacy protocols disabled), comprehensive HTTP security headers, and server version banner suppression.', bold_prefix='Module 2: ')
add_bullet('Database Security — PostgreSQL with least-privilege portal_user account, localhost-only connections via pg_hba.conf, encrypted password storage (bcrypt), and .env credential isolation excluded from version control.', bold_prefix='Module 3: ')
add_bullet('Application Security — CSRF protection, secure session cookies, WTForms input validation, RBAC enforcement, multi-layer file upload validation, rate limiting, Open Redirect remediation, dynamic IP blacklisting, custom error pages, and pinned dependency versions with CVE remediation.', bold_prefix='Module 4: ')
add_bullet('Security Monitoring — 20+ audit event types with IP tracking, admin log viewer dashboard, PDF anomaly report generator, Fail2ban automated intrusion response, Gunicorn structured logs, secure log rotation, and Rsyslog/Logwatch system-level logging.', bold_prefix='Module 5: ')
add_bullet('Business Continuity — Automated daily encrypted backup (GPG AES-256) with Cron scheduling, simulated crash script, automated restore script with RTO stopwatch measurement, 7-day retention policy, and documented MTD/MTO/RTO/RPO compliance parameters.', bold_prefix='Module 6: ')
add_body('')
add_body('The system was validated through automated end-to-end integration testing and security scanning. All identified vulnerabilities (Snyk Code CWE-601, GitHub CodeQL URL Redirection, CVE-2024-34069 Werkzeug RCE) were successfully remediated.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. REFERENCES
# ═══════════════════════════════════════════════════════════════
add_heading_styled('9. REFERENCES', level=1)
refs = [
    '[1] Flask Documentation, https://flask.palletsprojects.com/en/3.0.x/',
    '[2] PostgreSQL 16 Documentation, https://www.postgresql.org/docs/16/',
    '[3] Nginx Security Hardening, https://docs.nginx.com/nginx/admin-guide/security-controls/',
    '[4] OWASP Application Security Verification Standard (ASVS), https://owasp.org/www-project-application-security-verification-standard/',
    '[5] CWE-601: URL Redirection to Untrusted Site, https://cwe.mitre.org/data/definitions/601.html',
    '[6] CVE-2024-34069: Werkzeug Remote Code Execution, https://nvd.nist.gov/vuln/detail/CVE-2024-34069',
    '[7] Flask-Limiter Documentation, https://flask-limiter.readthedocs.io/en/stable/',
    '[8] Fail2ban Documentation, https://www.fail2ban.org/wiki/index.php/Main_Page',
    '[9] UFW Documentation, https://help.ubuntu.com/community/UFW',
    '[10] GPG Symmetric Encryption, https://www.gnupg.org/gph/en/manual/x110.html',
    '[11] PAM pam_pwquality Module, https://linux.die.net/man/8/pam_pwquality',
    '[12] Snyk Code – SAST, https://snyk.io/product/snyk-code/',
    '[13] GitHub CodeQL Analysis, https://codeql.github.com/',
    '[14] Gunicorn Documentation, https://docs.gunicorn.org/en/stable/',
    '[15] Flask-Bcrypt Password Hashing, https://flask-bcrypt.readthedocs.io/en/1.0.1/',
]
for ref in refs:
    add_body(ref)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# APPENDIX B – SYSTEMD SERVICE
# ═══════════════════════════════════════════════════════════════
add_heading_styled('APPENDIX A – GUNICORN SYSTEMD SERVICE UNIT', level=1)
add_code_block("""[Unit]
Description=Gunicorn instance to serve UniKL PJ Research Collaboration Portal
After=network.target

[Service]
User=jurus
Group=www-data
WorkingDirectory=/home/jurus/Documents/Project/UniPortal-Jurus
Environment="PATH=.../venv/bin:/usr/bin:/usr/local/bin"
Environment="FLASK_ENV=production"
ExecStart=.../venv/bin/gunicorn --workers 3 --bind 127.0.0.1:5000 \\
    --access-logfile /var/log/uniklpj/access.log \\
    --error-logfile /var/log/uniklpj/error.log \\
    wsgi:app
Restart=always

[Install]
WantedBy=multi-user.target""")

add_heading_styled('APPENDIX B – ENVIRONMENT TEMPLATE (.env.example)', level=1)
add_code_block("""# Flask Settings
SECRET_KEY=generate_a_random_secure_hex_key_here
FLASK_ENV=production
FLASK_DEBUG=False

# Database configuration
DATABASE_URL=postgresql://portal_user:secure_password_here@localhost/uniklpj_db

# File Upload Settings
UPLOAD_FOLDER=/var/www/uniklpj_uploads/""")

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
output_path = '/home/jurus/Documents/Project/UniPortal-Jurus/Jurus_FarhanIsmail.docx'
doc.save(output_path)
print(f"[+] Report generated successfully: {output_path}")
print(f"[+] File size: {os.path.getsize(output_path)} bytes")
